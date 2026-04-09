#!/usr/bin/env python3
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


@dataclass
class Block:
    kind: str
    data: object


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def strip_bib_markup(text: str) -> str:
    text = text.replace("{", "").replace("}", "")
    text = text.replace(r"\%", "%")
    text = text.replace(r"\&", "&")
    text = text.replace(r"~", " ")
    text = text.replace("\\\\", " ")
    text = text.replace("--", "-")
    return " ".join(text.split())


def format_authors(author_text: str) -> str:
    parts = [part.strip() for part in author_text.split(" and ") if part.strip()]
    if not parts:
        return author_text
    if len(parts) == 1:
        return parts[0]
    if len(parts) == 2:
        return f"{parts[0]} and {parts[1]}"
    return ", ".join(parts[:-1]) + f", and {parts[-1]}"


def clean_inline_latex(text: str, citation_map: Dict[str, int]) -> str:
    def cite_repl(match: re.Match[str]) -> str:
        keys = [part.strip() for part in match.group(1).split(",") if part.strip()]
        nums = [str(citation_map[key]) for key in keys if key in citation_map]
        return "[" + ", ".join(nums) + "]" if nums else ""

    replacements = [
        (r"~", " "),
        (r"\\%", "%"),
        (r"\\texttt\{([^}]*)\}", r"\1"),
        (r"\\textbf\{([^}]*)\}", r"\1"),
        (r"\\textit\{([^}]*)\}", r"\1"),
        (r"\\url\{([^}]*)\}", r"\1"),
        (r"\\label\{[^}]*\}", ""),
        (r"\\ref\{([^}]*)\}", r"\1"),
    ]

    text = re.sub(r"\\cite\{([^}]*)\}", cite_repl, text)
    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text)
    return " ".join(text.split())


def format_display_equation(equation: str, citation_map: Dict[str, int]) -> str:
    equation = equation.strip()
    if equation == r"\eta_{\text{path}} = \frac{P_{\text{IT}}}{P_{\text{source}}}.":
        return "η_path = P_IT / P_source"

    equation = clean_inline_latex(equation, citation_map)
    equation = re.sub(r"\\text\{([^}]*)\}", r"\1", equation)
    equation = re.sub(r"\\frac\{([^}]*)\}\{([^}]*)\}", r"(\1) / (\2)", equation)
    equation = equation.replace(r"\eta", "η")
    equation = re.sub(r"_\{([^}]*)\}", r"_\1", equation)
    equation = equation.replace("{", "").replace("}", "")
    equation = equation.replace("\\", "")
    return " ".join(equation.split())


def ordered_citation_keys(tex: str) -> List[str]:
    keys: List[str] = []
    for match in re.finditer(r"\\cite\{([^}]*)\}", tex):
        for key in [part.strip() for part in match.group(1).split(",") if part.strip()]:
            if key not in keys:
                keys.append(key)
    return keys


def parse_bib_entries(bib_text: str) -> Dict[str, Dict[str, str]]:
    entries: Dict[str, Dict[str, str]] = {}
    chunks = re.split(r"(?=@\w+\{)", bib_text)
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        head_match = re.match(r"@(\w+)\{([^,]+),", chunk)
        if not head_match:
            continue
        key = head_match.group(2).strip()
        fields: Dict[str, str] = {"entry_type": head_match.group(1)}
        for field, value in re.findall(r"(\w+)\s*=\s*\{((?:[^{}]|\{[^{}]*\})*)\}", chunk, flags=re.S):
            fields[field.lower()] = " ".join(value.replace("\n", " ").split())
        entries[key] = fields
    return entries


def format_reference(entry: Dict[str, str], number: int) -> str:
    author = format_authors(strip_bib_markup(entry.get("author", "Unknown author")))
    title = strip_bib_markup(entry.get("title", "Untitled"))
    year = strip_bib_markup(entry.get("year", "n.d."))
    journal = strip_bib_markup(entry.get("journal", "")) or None
    booktitle = strip_bib_markup(entry.get("booktitle", "")) or None
    institution = strip_bib_markup(entry.get("institution", "")) or None
    volume = strip_bib_markup(entry.get("volume", "")) or None
    number_field = strip_bib_markup(entry.get("number", "")) or None
    pages = strip_bib_markup(entry.get("pages", "")) or None
    doi = strip_bib_markup(entry.get("doi", "")) or None
    url = strip_bib_markup(entry.get("url", "")) or None

    parts = [f"[{number}] {author}, \"{title}.\""]
    if journal:
        line = journal
        if volume:
            line += f", vol. {volume}"
        if number_field:
            line += f", no. {number_field}"
        if pages:
            line += f", pp. {pages}"
        line += f", {year}."
        parts.append(line)
    elif booktitle:
        line = f"In {booktitle}"
        if pages:
            line += f", pp. {pages}"
        line += f", {year}."
        parts.append(line)
    elif institution:
        parts.append(f"{institution}, {year}.")
    else:
        parts.append(f"{year}.")
    if doi:
        parts.append(f"doi: {doi}.")
    elif url:
        parts.append(url)
    return " ".join(parts)


def extract_environment(text: str, env_name: str) -> str:
    match = re.search(rf"\\begin\{{{env_name}\}}(.*?)\\end\{{{env_name}\}}", text, flags=re.S)
    return match.group(1).strip() if match else ""


def parse_table_block(block_text: str, citation_map: Dict[str, int]) -> Dict[str, object]:
    caption_match = re.search(r"\\caption\{(.*?)\}", block_text, flags=re.S)
    caption = clean_inline_latex(caption_match.group(1), citation_map) if caption_match else "Table"
    label_match = re.search(r"\\label\{(.*?)\}", block_text)
    label = label_match.group(1) if label_match else ""
    tabular_match = re.search(r"\\begin\{tabular\}\{.*?\}(.*?)\\end\{tabular\}", block_text, flags=re.S)
    raw_tabular = tabular_match.group(1) if tabular_match else ""
    raw_lines = [line.strip() for line in raw_tabular.splitlines() if line.strip()]
    rows: List[List[str]] = []
    for raw_line in raw_lines:
        if raw_line.startswith("\\toprule") or raw_line.startswith("\\midrule") or raw_line.startswith("\\bottomrule"):
            continue
        if "\\\\" not in raw_line:
            continue
        row_text = raw_line.split("\\\\")[0].strip()
        columns = [clean_inline_latex(part.strip(), citation_map) for part in row_text.split("&")]
        rows.append(columns)
    headers = rows[0] if rows else []
    body = rows[1:] if len(rows) > 1 else []
    return {"caption": caption, "label": label, "headers": headers, "rows": body}


def parse_figure_block(block_text: str, citation_map: Dict[str, int], base_dir: Path) -> Dict[str, object]:
    path_match = re.search(r"\\includegraphics\[.*?\]\{(.*?)\}", block_text)
    caption_match = re.search(r"\\caption\{(.*?)\}", block_text, flags=re.S)
    label_match = re.search(r"\\label\{(.*?)\}", block_text)
    path = base_dir / path_match.group(1) if path_match else None
    caption = clean_inline_latex(caption_match.group(1), citation_map) if caption_match else ""
    label = label_match.group(1) if label_match else ""
    return {"path": path, "caption": caption, "label": label}


def collect_blocks(tex: str, citation_map: Dict[str, int], base_dir: Path) -> List[Block]:
    body_match = re.search(r"\\begin\{document\}(.*?)\\bibliographystyle", tex, flags=re.S)
    body = body_match.group(1) if body_match else tex
    body = re.sub(r"\\maketitle", "", body)
    body = re.sub(r"\\tableofcontents", "", body)

    blocks: List[Block] = []
    pos = 0
    pattern = re.compile(
        r"\\begin\{abstract\}.*?\\end\{abstract\}|"
        r"\\section\*?\{.*?\}|"
        r"\\subsection\{.*?\}|"
        r"\\begin\{figure\}.*?\\end\{figure\}|"
        r"\\begin\{table\}.*?\\end\{table\}|"
        r"\\begin\{itemize\}.*?\\end\{itemize\}|"
        r"\\begin\{equation\}.*?\\end\{equation\}|"
        r"\\begin\{verbatim\}.*?\\end\{verbatim\}",
        flags=re.S,
    )

    for match in pattern.finditer(body):
        preceding = body[pos:match.start()].strip()
        if preceding:
            for paragraph in [part.strip() for part in re.split(r"\n\s*\n", preceding) if part.strip()]:
                blocks.append(Block("paragraph", clean_inline_latex(paragraph, citation_map)))
        token = match.group(0)
        if token.startswith("\\begin{abstract}"):
            blocks.append(Block("heading1", "Abstract"))
            blocks.append(Block("paragraph", clean_inline_latex(extract_environment(token, "abstract"), citation_map)))
        elif token.startswith("\\section"):
            title = re.search(r"\\section\*?\{(.*?)\}", token, flags=re.S).group(1)
            blocks.append(Block("heading1", clean_inline_latex(title, citation_map)))
        elif token.startswith("\\subsection"):
            title = re.search(r"\\subsection\{(.*?)\}", token, flags=re.S).group(1)
            blocks.append(Block("heading2", clean_inline_latex(title, citation_map)))
        elif token.startswith("\\begin{figure}"):
            blocks.append(Block("figure", parse_figure_block(token, citation_map, base_dir)))
        elif token.startswith("\\begin{table}"):
            blocks.append(Block("table", parse_table_block(token, citation_map)))
        elif token.startswith("\\begin{itemize}"):
            raw_items = re.findall(r"\\item\s+(.*?)(?=(?:\\item|\\end\{itemize\}))", token, flags=re.S)
            items = [clean_inline_latex(item.strip(), citation_map) for item in raw_items if item.strip()]
            blocks.append(Block("bullet_list", items))
        elif token.startswith("\\begin{equation}"):
            equation = extract_environment(token, "equation").strip()
            equation = format_display_equation(equation, citation_map)
            blocks.append(Block("equation", equation))
        elif token.startswith("\\begin{verbatim}"):
            blocks.append(Block("code", extract_environment(token, "verbatim")))
        pos = match.end()

    trailing = body[pos:].strip()
    if trailing:
        for paragraph in [part.strip() for part in re.split(r"\n\s*\n", trailing) if part.strip()]:
            if paragraph:
                blocks.append(Block("paragraph", clean_inline_latex(paragraph, citation_map)))
    return blocks


def build_docx(output_path: Path, title: str, author: str, blocks: List[Block], references: List[str]) -> None:
    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author)
    run.font.size = Pt(11)

    for block in blocks:
        if block.kind == "heading1":
            document.add_heading(block.data, level=1)
        elif block.kind == "heading2":
            document.add_heading(block.data, level=2)
        elif block.kind == "paragraph":
            document.add_paragraph(str(block.data))
        elif block.kind == "bullet_list":
            for item in block.data:
                document.add_paragraph(item, style="List Bullet")
        elif block.kind == "equation":
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(block.data))
            run.italic = True
        elif block.kind == "code":
            p = document.add_paragraph()
            run = p.add_run(str(block.data))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif block.kind == "figure":
            figure = block.data
            if figure["path"] and Path(figure["path"]).exists():
                document.add_picture(str(figure["path"]), width=Inches(6.0))
                cap = document.add_paragraph(figure["caption"])
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif block.kind == "table":
            table_data = block.data
            document.add_paragraph(table_data["caption"])
            headers = table_data["headers"]
            rows = table_data["rows"]
            if headers:
                table = document.add_table(rows=len(rows) + 1, cols=len(headers))
                table.style = "Table Grid"
                for idx, header in enumerate(headers):
                    table.cell(0, idx).text = header
                for row_idx, row in enumerate(rows, start=1):
                    for col_idx, value in enumerate(row):
                        table.cell(row_idx, col_idx).text = value
            document.add_paragraph("")

    document.add_heading("References", level=1)
    for reference in references:
        document.add_paragraph(reference)
    document.save(str(output_path))


def fit_image(path: Path, max_width: float, max_height: float) -> tuple[float, float]:
    from PIL import Image as PILImage

    with PILImage.open(path) as img:
        width, height = img.size
    scale = min(max_width / width, max_height / height, 1.0)
    return width * scale, height * scale


def build_pdf(output_path: Path, title: str, author: str, blocks: List[Block], references: List[str]) -> None:
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="PaperTitle", parent=styles["Title"], alignment=TA_CENTER, fontName="Times-Bold", fontSize=16, leading=20, spaceAfter=8))
    styles.add(ParagraphStyle(name="PaperAuthor", parent=styles["Normal"], alignment=TA_CENTER, fontName="Times-Roman", fontSize=11, spaceAfter=14))
    styles.add(ParagraphStyle(name="PaperBody", parent=styles["BodyText"], fontName="Times-Roman", fontSize=10.5, leading=14, spaceAfter=8))
    styles.add(ParagraphStyle(name="PaperHeading1", parent=styles["Heading1"], fontName="Times-Bold", fontSize=14, leading=18, spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="PaperHeading2", parent=styles["Heading2"], fontName="Times-BoldItalic", fontSize=11.5, leading=14, spaceBefore=8, spaceAfter=4))
    styles.add(ParagraphStyle(name="Caption", parent=styles["Italic"], alignment=TA_CENTER, fontName="Times-Italic", fontSize=9.5, leading=12, spaceAfter=8))
    styles.add(ParagraphStyle(name="Equation", parent=styles["BodyText"], alignment=TA_CENTER, fontName="Times-Italic", fontSize=10.5, leading=14, spaceAfter=8))
    styles.add(ParagraphStyle(name="Reference", parent=styles["BodyText"], fontName="Times-Roman", fontSize=9.5, leading=12, leftIndent=0, firstLineIndent=0, spaceAfter=4))
    styles.add(ParagraphStyle(name="CodeBlock", parent=styles["BodyText"], fontName="Courier", fontSize=8.5, leading=10, spaceAfter=8))

    story = [
        Paragraph(title, styles["PaperTitle"]),
        Paragraph(author, styles["PaperAuthor"]),
    ]

    for block in blocks:
        if block.kind == "heading1":
            story.append(Paragraph(str(block.data), styles["PaperHeading1"]))
        elif block.kind == "heading2":
            story.append(Paragraph(str(block.data), styles["PaperHeading2"]))
        elif block.kind == "paragraph":
            story.append(Paragraph(str(block.data), styles["PaperBody"]))
        elif block.kind == "bullet_list":
            items = [ListItem(Paragraph(item, styles["PaperBody"])) for item in block.data]
            story.append(ListFlowable(items, bulletType="bullet", start="circle", leftIndent=18))
            story.append(Spacer(1, 0.08 * inch))
        elif block.kind == "equation":
            story.append(Paragraph(str(block.data), styles["Equation"]))
        elif block.kind == "code":
            story.append(Preformatted(str(block.data), styles["CodeBlock"]))
        elif block.kind == "figure":
            figure = block.data
            if figure["path"] and Path(figure["path"]).exists():
                width, height = fit_image(Path(figure["path"]), max_width=6.2 * inch, max_height=4.7 * inch)
                story.append(Image(str(figure["path"]), width=width, height=height))
                story.append(Paragraph(figure["caption"], styles["Caption"]))
        elif block.kind == "table":
            table_data = block.data
            story.append(Paragraph(table_data["caption"], styles["Caption"]))
            rows = [table_data["headers"]] + table_data["rows"]
            if rows and rows[0]:
                reportlab_table = Table(rows, repeatRows=1)
                reportlab_table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                            ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
                            ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
                            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                            ("LEADING", (0, 0), (-1, -1), 10),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                    )
                )
                story.append(reportlab_table)
                story.append(Spacer(1, 0.12 * inch))

    story.append(PageBreak())
    story.append(Paragraph("References", styles["PaperHeading1"]))
    for reference in references:
        story.append(Paragraph(reference, styles["Reference"]))

    doc = SimpleDocTemplate(str(output_path), pagesize=letter, rightMargin=0.8 * inch, leftMargin=0.8 * inch, topMargin=0.8 * inch, bottomMargin=0.8 * inch)
    doc.build(story)


def main() -> None:
    base_dir = Path(__file__).resolve().parent
    tex_path = base_dir / "dc_subtransmission_backbone_position_paper.tex"
    bib_path = base_dir / "references.bib"
    tex = read_text(tex_path)
    bib = read_text(bib_path)

    citation_keys = ordered_citation_keys(tex)
    citation_map = {key: idx + 1 for idx, key in enumerate(citation_keys)}
    bib_entries = parse_bib_entries(bib)
    references = [format_reference(bib_entries[key], citation_map[key]) for key in citation_keys if key in bib_entries]

    title_match = re.search(r"\\title\{(.*?)\}", tex, flags=re.S)
    author_match = re.search(r"\\author\{(.*?)\}", tex, flags=re.S)
    title = clean_inline_latex(title_match.group(1).replace("\\\\", ": "), citation_map) if title_match else "White Paper"
    author = clean_inline_latex(author_match.group(1), citation_map) if author_match else ""

    blocks = collect_blocks(tex, citation_map, base_dir)

    build_docx(base_dir / "dc_subtransmission_backbone_position_paper.docx", title, author, blocks, references)
    build_pdf(base_dir / "dc_subtransmission_backbone_position_paper.pdf", title, author, blocks, references)
    print("Wrote DOCX and PDF outputs.")


if __name__ == "__main__":
    main()
